"""
KnowledgeOrganizerAgent - Intelligent Document Organization

This agent:
1. Explores a folder of knowledge files (PDF, TXT, DOCX, MD, CSV, XLSX)
2. Extracts table of contents from PDFs when available
3. Chunks documents based on TOC structure or intelligent reasoning
4. Creates organized folder structure mirroring document organization
5. Saves chunks in a structured "knowledge-files-organized" directory

Chunker Tools (leveraging established libraries):
- PDFChunker: TOC-based hierarchical splitting with bookmark extraction
- MarkdownChunker: Uses langchain MarkdownHeaderTextSplitter for header-based splitting
- CSVChunker: Uses pandas for efficient row-based chunking
- ExcelChunker: Uses pandas + openpyxl for sheet-aware processing

Author: Reza Rahimi
Date: December 20, 2025
"""

import os
import sys
import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from abc import ABC, abstractmethod
import PyPDF2
import time

# OCR for scanned PDFs
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Add project root to Python path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.prompt_manager import get_prompt_manager
from utils.llm_client import create_llm_client
from utils.config import get_config

# Optional: langchain text splitters for Markdown
try:
    from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
    LANGCHAIN_SUPPORT = True
except ImportError:
    LANGCHAIN_SUPPORT = False
    print("Note: langchain-text-splitters not installed. Markdown chunking will use fallback. Install with: pip install langchain-text-splitters")

# Optional: pandas for CSV/Excel
try:
    import pandas as pd
    PANDAS_SUPPORT = True
except ImportError:
    PANDAS_SUPPORT = False
    print("Note: pandas not installed. CSV/Excel chunking disabled. Install with: pip install pandas openpyxl")

# Optional: openpyxl for Excel (used by pandas)
try:
    import openpyxl
    EXCEL_SUPPORT = PANDAS_SUPPORT  # Excel requires both pandas and openpyxl
except ImportError:
    EXCEL_SUPPORT = False
    if PANDAS_SUPPORT:
        print("Note: openpyxl not installed. Excel support disabled. Install with: pip install openpyxl")

# Optional: python-docx for Word documents
try:
    from docx import Document as DocxDocument
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Note: python-docx not installed. DOCX support disabled. Install with: pip install python-docx")


@dataclass
class DocumentChunk:
    """Represents a chunk of a document."""
    chunk_id: str
    title: str
    content: str
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    row_start: Optional[int] = None  # For CSV/Excel
    row_end: Optional[int] = None    # For CSV/Excel
    sheet_name: Optional[str] = None  # For Excel
    section_path: List[str] = None  # Hierarchical path (e.g., ["Chapter 1", "Section 1.1"])
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.section_path is None:
            self.section_path = []
        if self.metadata is None:
            self.metadata = {}


@dataclass
class TableOfContentsEntry:
    """Represents a TOC entry."""
    title: str
    page: Optional[int] = None
    level: int = 0  # Hierarchy level (0=chapter, 1=section, 2=subsection)
    children: List['TableOfContentsEntry'] = None
    
    def __post_init__(self):
        if self.children is None:
            self.children = []


# =============================================================================
# CHUNKER TOOLS - Base Class and Implementations
# =============================================================================

class BaseChunker(ABC):
    """
    Abstract base class for document chunkers.
    
    Each chunker implements format-specific logic for:
    - Detecting if it can handle a file
    - Extracting structured chunks with metadata
    """
    
    @abstractmethod
    def can_handle(self, file_path: Path) -> bool:
        """Check if this chunker can handle the given file."""
        pass
    
    @abstractmethod
    def chunk(self, file_path: Path, config: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Chunk the document into structured pieces."""
        pass
    
    @property
    @abstractmethod
    def supported_extensions(self) -> set:
        """Return set of supported file extensions."""
        pass
    
    @property
    @abstractmethod
    def name(self) -> str:
        """Return the chunker name."""
        pass


class MarkdownChunker(BaseChunker):
    """
    Markdown Chunker Tool
    
    Uses langchain-text-splitters MarkdownHeaderTextSplitter for intelligent
    header-based splitting that preserves hierarchy metadata.
    Falls back to regex-based splitting if langchain not available.
    """
    
    @property
    def name(self) -> str:
        return "MarkdownChunker (langchain)" if LANGCHAIN_SUPPORT else "MarkdownChunker (fallback)"
    
    @property
    def supported_extensions(self) -> set:
        return {'.md', '.markdown'}
    
    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.supported_extensions
    
    def chunk(self, file_path: Path, config: Dict[str, Any] = None) -> List[DocumentChunk]:
        """
        Chunk markdown by headers using langchain MarkdownHeaderTextSplitter.
        
        Args:
            file_path: Path to markdown file
            config: Optional config with:
                - chunk_size: Max chunk size in chars for secondary splitting (default: 10000)
                - chunk_overlap: Overlap between chunks (default: 200)
        
        Returns:
            List of DocumentChunk objects
        """
        config = config or {}
        # Default ~2000 words ≈ 10000 chars; config.json max_chunk_size is in words
        chunk_size = config.get('chunk_size', 10000)
        chunk_overlap = config.get('chunk_overlap', 200)
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        except Exception as e:
            print(f"  ✗ Error reading markdown file: {e}")
            return []
        
        if LANGCHAIN_SUPPORT:
            return self._chunk_with_langchain(file_path, content, chunk_size, chunk_overlap)
        else:
            return self._chunk_with_fallback(file_path, content)
    
    def _chunk_with_langchain(self, file_path: Path, content: str, 
                               chunk_size: int, chunk_overlap: int) -> List[DocumentChunk]:
        """Use langchain MarkdownHeaderTextSplitter for intelligent splitting."""
        print(f"  → Using langchain MarkdownHeaderTextSplitter...")
        
        # Define headers to split on (H1-H4)
        headers_to_split_on = [
            ("#", "h1"),
            ("##", "h2"),
            ("###", "h3"),
            ("####", "h4"),
        ]
        
        try:
            # First split by headers
            md_splitter = MarkdownHeaderTextSplitter(
                headers_to_split_on=headers_to_split_on,
                strip_headers=False  # Keep headers in content
            )
            md_header_splits = md_splitter.split_text(content)
            
            # If chunks are too large, split further with RecursiveCharacterTextSplitter
            if any(len(doc.page_content) > chunk_size for doc in md_header_splits):
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    separators=["\n\n", "\n", " ", ""]
                )
                final_docs = text_splitter.split_documents(md_header_splits)
            else:
                final_docs = md_header_splits
            
            chunks = []
            for i, doc in enumerate(final_docs):
                # Build section path from metadata
                section_path = []
                for level in ['h1', 'h2', 'h3', 'h4']:
                    if level in doc.metadata:
                        section_path.append(doc.metadata[level])
                
                if not section_path:
                    section_path = [file_path.stem]
                
                title = section_path[-1] if section_path else f"Section {i+1}"
                
                chunk = DocumentChunk(
                    chunk_id=f"{file_path.stem}_{i+1:03d}",
                    title=title,
                    content=doc.page_content.strip(),
                    section_path=section_path,
                    metadata={
                        "source_file": file_path.name,
                        "chunk_method": "langchain_markdown_headers",
                        "format": "markdown",
                        "header_metadata": doc.metadata,
                        "has_code_blocks": '```' in doc.page_content
                    }
                )
                chunks.append(chunk)
            
            print(f"  ✓ Created {len(chunks)} chunks using langchain MarkdownHeaderTextSplitter")
            return chunks
            
        except Exception as e:
            print(f"  ⚠ Langchain splitting failed: {e}, falling back to regex")
            return self._chunk_with_fallback(file_path, content)
    
    def _chunk_with_fallback(self, file_path: Path, content: str) -> List[DocumentChunk]:
        """Fallback regex-based markdown splitting."""
        print(f"  → Using fallback regex-based markdown splitting...")
        
        chunks = []
        header_pattern = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        headers = []
        for match in header_pattern.finditer(content):
            headers.append({
                'level': len(match.group(1)),
                'title': match.group(2).strip(),
                'start': match.start(),
                'end': match.end()
            })
        
        if not headers:
            chunk = DocumentChunk(
                chunk_id=f"{file_path.stem}_001",
                title=file_path.stem,
                content=content.strip(),
                section_path=[file_path.stem],
                metadata={
                    "source_file": file_path.name,
                    "chunk_method": "markdown_no_headers",
                    "format": "markdown"
                }
            )
            return [chunk]
        
        section_stack = []
        for i, header in enumerate(headers):
            content_end = headers[i + 1]['start'] if i + 1 < len(headers) else len(content)
            chunk_content = content[header['start']:content_end].strip()
            
            while section_stack and section_stack[-1][0] >= header['level']:
                section_stack.pop()
            section_stack.append((header['level'], header['title']))
            
            chunk = DocumentChunk(
                chunk_id=f"{file_path.stem}_{i+1:03d}",
                title=header['title'],
                content=chunk_content,
                section_path=[s[1] for s in section_stack],
                metadata={
                    "source_file": file_path.name,
                    "chunk_method": "markdown_regex_fallback",
                    "format": "markdown",
                    "header_level": header['level']
                }
            )
            chunks.append(chunk)
        
        print(f"  ✓ Created {len(chunks)} chunks using regex fallback")
        return chunks


class CSVChunker(BaseChunker):
    """
    CSV Chunker Tool
    
    Uses pandas for efficient CSV chunking with streaming support for large files.
    Falls back to built-in csv module if pandas not available.
    """
    
    @property
    def name(self) -> str:
        return "CSVChunker (pandas)" if PANDAS_SUPPORT else "CSVChunker (csv)"
    
    @property
    def supported_extensions(self) -> set:
        return {'.csv', '.tsv'}
    
    def can_handle(self, file_path: Path) -> bool:
        if not PANDAS_SUPPORT:
            return False
        return file_path.suffix.lower() in self.supported_extensions
    
    def chunk(self, file_path: Path, config: Dict[str, Any] = None) -> List[DocumentChunk]:
        """
        Chunk CSV using pandas for efficient processing.
        
        Args:
            file_path: Path to CSV file
            config: Optional config with:
                - rows_per_chunk: Number of rows per chunk (default: 50)
                - include_header: Include column names context (default: True)
        
        Returns:
            List of DocumentChunk objects
        """
        if not PANDAS_SUPPORT:
            print(f"  ✗ pandas not installed. Cannot process CSV files.")
            return []
        
        config = config or {}
        rows_per_chunk = config.get('rows_per_chunk', 50)
        include_header = config.get('include_header', True)
        
        print(f"  → Using pandas to chunk CSV ({rows_per_chunk} rows per chunk)...")
        
        try:
            # Detect separator
            sep = '\t' if file_path.suffix.lower() == '.tsv' else ','
            
            # Read CSV with pandas (more robust than csv module)
            df = pd.read_csv(file_path, sep=sep, encoding='utf-8', on_bad_lines='skip')
            
            chunks = []
            total_rows = len(df)
            columns = list(df.columns)
            
            chunk_num = 1
            for start_idx in range(0, total_rows, rows_per_chunk):
                end_idx = min(start_idx + rows_per_chunk, total_rows)
                chunk_df = df.iloc[start_idx:end_idx]
                
                # Format content with column context
                content_lines = []
                if include_header:
                    content_lines.append(f"Columns: {', '.join(columns)}")
                    content_lines.append("")
                
                # Convert chunk to readable format
                for idx, row in chunk_df.iterrows():
                    row_parts = []
                    for col in columns:
                        val = row[col]
                        if pd.notna(val):
                            row_parts.append(f"{col}: {val}")
                    content_lines.append(f"Row {idx + 1}: {' | '.join(row_parts)}")
                
                row_start = start_idx + 1
                row_end = end_idx
                
                chunk = DocumentChunk(
                    chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                    title=f"Rows {row_start}-{row_end}",
                    content='\n'.join(content_lines),
                    row_start=row_start,
                    row_end=row_end,
                    section_path=[file_path.stem, f"Rows {row_start}-{row_end}"],
                    metadata={
                        "source_file": file_path.name,
                        "chunk_method": "pandas_csv",
                        "format": "csv",
                        "columns": columns,
                        "total_rows": total_rows,
                        "rows_in_chunk": len(chunk_df)
                    }
                )
                chunks.append(chunk)
                chunk_num += 1
            
            print(f"  ✓ Created {len(chunks)} chunks from {total_rows} CSV rows using pandas")
            return chunks
            
        except Exception as e:
            print(f"  ✗ Error processing CSV with pandas: {e}")
            return []


class ExcelChunker(BaseChunker):
    """
    Excel Chunker Tool
    
    Uses pandas + openpyxl for efficient Excel processing with sheet-aware chunking.
    Handles multiple sheets and large files efficiently.
    """
    
    @property
    def name(self) -> str:
        return "ExcelChunker (pandas)"
    
    @property
    def supported_extensions(self) -> set:
        return {'.xlsx', '.xls', '.xlsm'}
    
    def can_handle(self, file_path: Path) -> bool:
        if not EXCEL_SUPPORT:
            return False
        return file_path.suffix.lower() in self.supported_extensions
    
    def chunk(self, file_path: Path, config: Dict[str, Any] = None) -> List[DocumentChunk]:
        """
        Chunk Excel using pandas for efficient multi-sheet processing.
        
        Args:
            file_path: Path to Excel file
            config: Optional config with:
                - rows_per_chunk: Number of rows per chunk (default: 50)
                - sheet_names: Specific sheets to process (default: None = all)
                - include_empty_sheets: Include sheets with no data (default: False)
        
        Returns:
            List of DocumentChunk objects
        """
        if not EXCEL_SUPPORT:
            print(f"  ✗ Excel support not available. Install: pip install pandas openpyxl")
            return []
        
        config = config or {}
        rows_per_chunk = config.get('rows_per_chunk', 50)
        target_sheets = config.get('sheet_names', None)
        include_empty_sheets = config.get('include_empty_sheets', False)
        
        print(f"  → Using pandas to chunk Excel ({rows_per_chunk} rows per chunk)...")
        
        try:
            # Read all sheets with pandas
            excel_file = pd.ExcelFile(file_path, engine='openpyxl')
            sheet_names = target_sheets if target_sheets else excel_file.sheet_names
            
            chunks = []
            chunk_num = 1
            
            for sheet_name in sheet_names:
                if sheet_name not in excel_file.sheet_names:
                    print(f"    ℹ Sheet '{sheet_name}' not found, skipping")
                    continue
                
                # Read sheet
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                if df.empty:
                    if include_empty_sheets:
                        chunk = DocumentChunk(
                            chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                            title=f"Sheet: {sheet_name} (empty)",
                            content=f"Sheet '{sheet_name}' contains no data.",
                            sheet_name=sheet_name,
                            section_path=[file_path.stem, sheet_name],
                            metadata={
                                "source_file": file_path.name,
                                "chunk_method": "pandas_excel",
                                "format": "excel",
                                "sheet_name": sheet_name,
                                "is_empty": True
                            }
                        )
                        chunks.append(chunk)
                        chunk_num += 1
                    continue
                
                columns = list(df.columns)
                total_rows = len(df)
                
                # Chunk the sheet data
                for start_idx in range(0, total_rows, rows_per_chunk):
                    end_idx = min(start_idx + rows_per_chunk, total_rows)
                    chunk_df = df.iloc[start_idx:end_idx]
                    
                    # Format content
                    content_lines = [f"Sheet: {sheet_name}", ""]
                    content_lines.append(f"Columns: {', '.join(str(c) for c in columns)}")
                    content_lines.append("")
                    
                    for idx, row in chunk_df.iterrows():
                        row_parts = []
                        for col in columns:
                            val = row[col]
                            if pd.notna(val):
                                row_parts.append(f"{col}: {val}")
                        if row_parts:
                            content_lines.append(f"Row {idx + 2}: {' | '.join(row_parts)}")
                    
                    row_start = start_idx + 1
                    row_end = end_idx
                    
                    chunk = DocumentChunk(
                        chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                        title=f"{sheet_name} - Rows {row_start}-{row_end}",
                        content='\n'.join(content_lines),
                        row_start=row_start,
                        row_end=row_end,
                        sheet_name=sheet_name,
                        section_path=[file_path.stem, sheet_name, f"Rows {row_start}-{row_end}"],
                        metadata={
                            "source_file": file_path.name,
                            "chunk_method": "pandas_excel",
                            "format": "excel",
                            "sheet_name": sheet_name,
                            "columns": [str(c) for c in columns],
                            "total_rows": total_rows,
                            "rows_in_chunk": len(chunk_df)
                        }
                    )
                    chunks.append(chunk)
                    chunk_num += 1
            
            print(f"  ✓ Created {len(chunks)} chunks from {len(sheet_names)} Excel sheet(s) using pandas")
            return chunks
            
        except Exception as e:
            print(f"  ✗ Error processing Excel with pandas: {e}")
            return []


class DocxChunker(BaseChunker):
    """
    DOCX Chunker Tool
    
    Uses python-docx for intelligent Word document processing with:
    - Heading-based hierarchical splitting (Heading 1-6 styles)
    - Table extraction with cell content preservation
    - Paragraph grouping with style awareness
    """
    
    @property
    def name(self) -> str:
        return "DocxChunker (python-docx)" if DOCX_SUPPORT else "DocxChunker (disabled)"
    
    @property
    def supported_extensions(self) -> set:
        return {'.docx'}
    
    def can_handle(self, file_path: Path) -> bool:
        if not DOCX_SUPPORT:
            return False
        return file_path.suffix.lower() in self.supported_extensions
    
    def chunk(self, file_path: Path, config: Dict[str, Any] = None) -> List[DocumentChunk]:
        """
        Chunk DOCX using python-docx for structure-aware processing.
        
        Args:
            file_path: Path to DOCX file
            config: Optional config with:
                - chunk_by_heading: Split by headings (default: True)
                - min_chunk_size: Minimum chars per chunk (default: 100)
                - include_tables: Include table content (default: True)
        
        Returns:
            List of DocumentChunk objects
        """
        if not DOCX_SUPPORT:
            print(f"  ✗ python-docx not installed. Cannot process DOCX files.")
            return []
        
        config = config or {}
        chunk_by_heading = config.get('chunk_by_heading', True)
        min_chunk_size = config.get('min_chunk_size', 100)
        include_tables = config.get('include_tables', True)
        
        print(f"  → Using python-docx to chunk Word document...")
        
        try:
            doc = DocxDocument(file_path)
            
            if chunk_by_heading:
                return self._chunk_by_headings(file_path, doc, min_chunk_size, include_tables)
            else:
                return self._chunk_by_paragraphs(file_path, doc, min_chunk_size, include_tables)
                
        except Exception as e:
            print(f"  ✗ Error processing DOCX: {e}")
            return []
    
    def _is_heading(self, paragraph) -> Optional[int]:
        """
        Check if paragraph is a heading and return its level.
        
        Returns:
            Heading level (1-9) or None if not a heading
        """
        style_name = paragraph.style.name if paragraph.style else ""
        
        # Check for "Heading X" style
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading ', '').strip())
                return level
            except ValueError:
                pass
        
        # Check for "Title" style (treat as H1)
        if style_name == 'Title':
            return 1
        
        # Check for "Subtitle" style (treat as H2)
        if style_name == 'Subtitle':
            return 2
        
        return None
    
    def _extract_table_content(self, table) -> str:
        """Extract content from a table as formatted text."""
        rows_text = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                if row_idx == 0:
                    rows_text.append(f"| {' | '.join(cells)} |")
                    rows_text.append(f"| {' | '.join(['---'] * len(cells))} |")
                else:
                    rows_text.append(f"| {' | '.join(cells)} |")
        return '\n'.join(rows_text)
    
    def _chunk_by_headings(self, file_path: Path, doc, min_chunk_size: int, 
                           include_tables: bool) -> List[DocumentChunk]:
        """Split document by heading styles."""
        chunks = []
        current_chunk_content = []
        current_heading = None
        current_level = 0
        section_stack = []  # Track hierarchy: [(level, title), ...]
        chunk_num = 1
        
        # Build reverse-lookup maps once to avoid O(n²) scans per element
        elem_to_para = {p._element: p for p in doc.paragraphs}
        elem_to_table = {t._element: t for t in doc.tables}

        # Iterate through document body elements (paragraphs and tables)
        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                para = elem_to_para.get(element)

                if para is None:
                    continue

                heading_level = self._is_heading(para)

                if heading_level:
                    # Save previous chunk if it has content
                    if current_chunk_content and current_heading:
                        content = '\n\n'.join(current_chunk_content).strip()
                        if len(content) >= min_chunk_size:
                            chunk = self._create_chunk(
                                file_path, chunk_num, current_heading,
                                content, [s[1] for s in section_stack]
                            )
                            chunks.append(chunk)
                            chunk_num += 1

                    # Update section stack for hierarchy
                    while section_stack and section_stack[-1][0] >= heading_level:
                        section_stack.pop()
                    section_stack.append((heading_level, para.text.strip()))

                    # Start new chunk
                    current_heading = para.text.strip()
                    current_level = heading_level
                    current_chunk_content = [f"{'#' * heading_level} {current_heading}"]
                else:
                    # Regular paragraph
                    text = para.text.strip()
                    if text:
                        current_chunk_content.append(text)

            # Handle tables
            elif element.tag.endswith('tbl') and include_tables:
                table = elem_to_table.get(element)
                if table is not None:
                    table_content = self._extract_table_content(table)
                    if table_content:
                        current_chunk_content.append(f"\n{table_content}\n")
        
        # Save final chunk
        if current_chunk_content:
            content = '\n\n'.join(current_chunk_content).strip()
            if len(content) >= min_chunk_size or not chunks:
                heading = current_heading or file_path.stem
                chunk = self._create_chunk(
                    file_path, chunk_num, heading,
                    content, [s[1] for s in section_stack] if section_stack else [file_path.stem]
                )
                chunks.append(chunk)
        
        # If no headings found, fall back to paragraph chunking
        if not chunks:
            return self._chunk_by_paragraphs(file_path, doc, min_chunk_size, include_tables)
        
        print(f"  ✓ Created {len(chunks)} chunks by heading structure using python-docx")
        return chunks
    
    def _chunk_by_paragraphs(self, file_path: Path, doc, min_chunk_size: int,
                              include_tables: bool) -> List[DocumentChunk]:
        """Fallback: chunk by grouping paragraphs."""
        chunks = []
        current_content = []
        chunk_num = 1
        target_size = get_config().get_docx_fallback_chunk_size()  # Target chars per chunk
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current_content.append(text)
            
            # Check if we've accumulated enough content
            current_text = '\n\n'.join(current_content)
            if len(current_text) >= target_size:
                chunk = DocumentChunk(
                    chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                    title=f"Section {chunk_num}",
                    content=current_text,
                    section_path=[file_path.stem, f"Section {chunk_num}"],
                    metadata={
                        "source_file": file_path.name,
                        "chunk_method": "docx_paragraphs",
                        "format": "docx"
                    }
                )
                chunks.append(chunk)
                chunk_num += 1
                current_content = []
        
        # Handle remaining content
        if current_content:
            content = '\n\n'.join(current_content)
            if len(content) >= min_chunk_size or not chunks:
                chunk = DocumentChunk(
                    chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                    title=f"Section {chunk_num}" if chunks else file_path.stem,
                    content=content,
                    section_path=[file_path.stem, f"Section {chunk_num}"],
                    metadata={
                        "source_file": file_path.name,
                        "chunk_method": "docx_paragraphs",
                        "format": "docx"
                    }
                )
                chunks.append(chunk)
        
        print(f"  ✓ Created {len(chunks)} chunks by paragraph grouping using python-docx")
        return chunks
    
    def _create_chunk(self, file_path: Path, chunk_num: int, title: str,
                      content: str, section_path: List[str]) -> DocumentChunk:
        """Create a DocumentChunk with standard metadata."""
        return DocumentChunk(
            chunk_id=f"{file_path.stem}_{chunk_num:03d}",
            title=title,
            content=content,
            section_path=section_path,
            metadata={
                "source_file": file_path.name,
                "chunk_method": "docx_headings",
                "format": "docx",
                "has_tables": '|' in content
            }
        )


# =============================================================================
# CHUNKER TOOL REGISTRY
# =============================================================================

class ChunkerToolRegistry:
    """
    Registry for managing and selecting appropriate chunker tools.
    """
    
    def __init__(self):
        self._chunkers: List[BaseChunker] = []
    
    def register(self, chunker: BaseChunker):
        """Register a chunker tool."""
        self._chunkers.append(chunker)
        return self
    
    def get_chunker(self, file_path: Path) -> Optional[BaseChunker]:
        """Get appropriate chunker for a file."""
        for chunker in self._chunkers:
            if chunker.can_handle(file_path):
                return chunker
        return None
    
    def get_all_extensions(self) -> set:
        """Get all supported extensions from all chunkers."""
        extensions = set()
        for chunker in self._chunkers:
            extensions.update(chunker.supported_extensions)
        return extensions
    
    def list_chunkers(self) -> List[str]:
        """List all registered chunker names."""
        return [c.name for c in self._chunkers]


# =============================================================================
# DOCUMENT CHUNKING AGENT
# =============================================================================


class DocumentChunkingAgent:
    """
    Agent for intelligently organizing knowledge files into structured folders.
    
    Uses a tool-based architecture with specialized chunkers:
    - PDFChunker: TOC-based hierarchical splitting
    - MarkdownChunker: Header-based splitting (H1-H6)
    - CSVChunker: Row-based grouping with column context
    - ExcelChunker: Sheet-aware processing
    """
    
    def __init__(self, api_key: str, model: Optional[str] = None, reasoning_effort: Optional[str] = None):
        """
        Initialize the knowledge organizer agent.
        
        Args:
            api_key: API key for LLM provider
            model: Optional override for reasoning model
            reasoning_effort: Optional override for reasoning effort level (low/medium/high)
        """
        self.config = get_config()
        self.model = model or self.config.get_reasoning_model()
        self.reasoning_effort = reasoning_effort or self.config.get_reasoning_effort()
        self.client = create_llm_client(
            api_key=api_key,
            model=self.model,
            timeout=self.config.get_timeout(),
            max_retries=self.config.get_max_retries()
        )
        self.prompt_manager = get_prompt_manager()
        
        # Initialize chunker tool registry
        self.chunker_registry = ChunkerToolRegistry()
        self._register_chunker_tools()
        
        # Supported extensions from all registered chunkers + built-in
        self.supported_extensions = self.chunker_registry.get_all_extensions()
        self.supported_extensions.update({'.pdf', '.txt', '.text'})  # Built-in support
    
    def _register_chunker_tools(self):
        """Register all available chunker tools."""
        self.chunker_registry.register(MarkdownChunker())
        self.chunker_registry.register(CSVChunker())
        self.chunker_registry.register(ExcelChunker())
        self.chunker_registry.register(DocxChunker())
        
        print(f"Registered chunker tools: {', '.join(self.chunker_registry.list_chunkers())}")
    
    def get_chunker_for_file(self, file_path: Path) -> Optional[BaseChunker]:
        """
        Get the appropriate chunker tool for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Chunker instance or None if no specialized chunker available
        """
        return self.chunker_registry.get_chunker(file_path)
    
    def explore_input(self, input_path: str, only_files: List[str] = None) -> List[Path]:
        """
        Explore input (file or folder) and find all supported knowledge files.
        
        Args:
            input_path: Path to a file or folder containing knowledge files
            only_files: Optional list of specific filenames to process (filter)
            
        Returns:
            List of Path objects for supported files
        """
        path = Path(input_path)
        
        if not path.exists():
            raise ValueError(f"Path does not exist: {input_path}")
        
        files = []
        
        # Check if it's a single file or a directory
        if path.is_file():
            if path.suffix.lower() in self.supported_extensions:
                files.append(path)
                print(f"Processing single file: {path.name}")
            else:
                raise ValueError(f"Unsupported file format: {path.suffix}. Supported: {self.supported_extensions}")
        else:
            # It's a directory - explore recursively
            for file_path in path.rglob("*"):
                if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                    files.append(file_path)
            
            # Filter to only selected files if specified
            if only_files:
                only_names = set(only_files)
                files = [f for f in files if f.name in only_names]
                print(f"Found {len(files)} knowledge files (filtered from selection):")
            else:
                print(f"Found {len(files)} knowledge files:")
            for file_path in files:
                # Indicate which chunker will handle each file
                chunker = self.get_chunker_for_file(file_path)
                chunker_name = chunker.name if chunker else "PDFChunker/TextChunker"
                print(f"  - {file_path.name} ({file_path.suffix}) → {chunker_name}")
        
        return files
        
    def extract_pdf_toc(self, pdf_path: Path) -> Optional[List[TableOfContentsEntry]]:
        """
        Extract table of contents from a PDF file.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of TOC entries or None if no TOC found
        """
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Try to get outline (bookmarks)
                if pdf_reader.outline:
                    toc = self._parse_pdf_outline(pdf_reader.outline, pdf_reader)
                    if toc:
                        print(f"  ✓ Extracted {len(toc)} TOC entries from PDF outline")
                        return toc
                
                # If no outline, try to extract from first few pages
                print("  → No PDF outline found, analyzing first pages for TOC...")
                toc = self._extract_toc_from_text(pdf_reader)
                if toc:
                    print(f"  ✓ Detected {len(toc)} TOC entries from document text")
                    return toc
                
                print("  ℹ No table of contents detected")
                return None
                
        except Exception as e:
            print(f"  ✗ Error extracting TOC from PDF: {e}")
            return None
    
    def _parse_pdf_outline(self, outline: List, pdf_reader: PyPDF2.PdfReader, level: int = 0) -> List[TableOfContentsEntry]:
        """
        Parse PDF outline/bookmarks into TOC entries.
        
        Args:
            outline: PDF outline structure
            pdf_reader: PDF reader object
            level: Current hierarchy level
            
        Returns:
            List of TOC entries
        """
        toc_entries = []
        
        for item in outline:
            if isinstance(item, list):
                # Nested list - recurse
                toc_entries.extend(self._parse_pdf_outline(item, pdf_reader, level + 1))
            else:
                # Bookmark item
                title = item.title if hasattr(item, 'title') else str(item)
                page = None
                
                # Try to get page number
                try:
                    if hasattr(item, 'page'):
                        page = pdf_reader.get_destination_page_number(item) + 1
                except Exception:
                    pass
                
                entry = TableOfContentsEntry(
                    title=title,
                    page=page,
                    level=level
                )
                toc_entries.append(entry)
        
        return toc_entries
    
    def _extract_toc_from_text(self, pdf_reader: PyPDF2.PdfReader) -> Optional[List[TableOfContentsEntry]]:
        """
        Extract TOC from the text of first few pages using AI reasoning.
        
        Args:
            pdf_reader: PDF reader object
            
        Returns:
            List of TOC entries or None
        """
        # Extract text from first 5 pages (where TOC usually is)
        toc_text = ""
        max_pages = min(5, len(pdf_reader.pages))
        
        for i in range(max_pages):
            try:
                page_text = pdf_reader.pages[i].extract_text()
                toc_text += f"\n--- Page {i+1} ---\n{page_text}"
            except Exception:
                continue
        
        if not toc_text.strip():
            return None
        
        # Use AI to detect and parse TOC
        prompt = f"""Analyze this document excerpt and determine if it contains a Table of Contents.

DOCUMENT EXCERPT:
{toc_text[:4000]}

TASK:
1. Determine if this contains a table of contents
2. If yes, extract all TOC entries with their titles and page numbers
3. Identify the hierarchy level of each entry (0=chapter, 1=section, 2=subsection)

Respond with JSON:
{{
  "has_toc": true/false,
  "toc_entries": [
    {{
      "title": "Chapter or section title",
      "page": page_number or null,
      "level": 0-2
    }}
  ]
}}

If no TOC is found, return {{"has_toc": false, "toc_entries": []}}
"""

        try:
            response = self.client.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0.3,
                max_tokens=8192,
                reasoning_effort=self.reasoning_effort
            )
            
            result = json.loads(response.choices[0].message.content)
            
            if result.get('has_toc') and result.get('toc_entries'):
                toc_entries = [
                    TableOfContentsEntry(
                        title=entry['title'],
                        page=entry.get('page'),
                        level=entry.get('level', 0)
                    )
                    for entry in result['toc_entries']
                ]
                return toc_entries
            
        except Exception as e:
            print(f"  ✗ Error analyzing TOC with AI: {e}")
        
        return None
    
    def chunk_pdf_with_toc(self, pdf_path: Path, toc: List[TableOfContentsEntry]) -> List[DocumentChunk]:
        """
        Chunk a PDF based on its table of contents.
        
        Args:
            pdf_path: Path to PDF file
            toc: Table of contents entries
            
        Returns:
            List of document chunks
        """
        print(f"  → Chunking based on TOC structure...")
        
        chunks = []
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

                # Extract page text once and reuse it across TOC sections.
                # Large guide PDFs can have thousands of outline entries that
                # otherwise trigger repeated extract_text() calls for the same pages.
                page_text_cache = []
                needs_ocr = False
                for page_num in range(total_pages):
                    try:
                        page_text = pdf_reader.pages[page_num].extract_text() or ""
                    except Exception:
                        page_text = ""
                    page_text_cache.append(page_text.strip())
                    if not page_text.strip():
                        needs_ocr = True

                    if total_pages >= 250 and ((page_num + 1) % 250 == 0 or page_num + 1 == total_pages):
                        print(f"    ... extracted text from {page_num + 1}/{total_pages} pages")

                # If any pages had no text, attempt OCR on those pages
                if needs_ocr and OCR_AVAILABLE:
                    print(f"  → Some pages have no extractable text, running OCR on empty pages...")
                    for page_num in range(total_pages):
                        if not page_text_cache[page_num]:
                            ocr_text = self._ocr_pdf_page(pdf_path, page_num)
                            if ocr_text.strip():
                                page_text_cache[page_num] = ocr_text.strip()

                # Pre-resolve page=0 entries.  Some PDFs report page 0 for
                # TOC bookmarks whose destinations can't be resolved.  We
                # inherit the page number from the nearest preceding entry
                # that has a valid (>0) page so that we don't accidentally
                # slice from page 1 to the end of the document.
                resolved_pages: List[int] = []
                last_valid_page = 1
                for entry in toc:
                    p = entry.page
                    if p is not None and p > 0:
                        last_valid_page = p
                        resolved_pages.append(p)
                    else:
                        resolved_pages.append(last_valid_page)
                
                for i, toc_entry in enumerate(toc):
                    # Determine page range for this chunk
                    page_start = resolved_pages[i]
                    if page_start < 1 or page_start > total_pages:
                        continue
                    
                    # End page is the start of next entry or last page
                    if i + 1 < len(toc):
                        page_end = resolved_pages[i + 1] - 1
                        # When next entry starts on the same page, include at least that page
                        if page_end < page_start:
                            page_end = page_start
                    else:
                        page_end = total_pages

                    if page_end < page_start:
                        continue
                    
                    # Extract content for this chunk
                    page_slice = page_text_cache[page_start - 1:min(page_end, total_pages)]
                    content = "\n\n".join(text for text in page_slice if text)
                    
                    if content.strip():
                        # Create section path based on TOC hierarchy
                        section_path = self._build_section_path(toc_entry, toc, i)
                        
                        chunk = DocumentChunk(
                            chunk_id=f"{pdf_path.stem}_{i+1:03d}",
                            title=toc_entry.title,
                            content=content.strip(),
                            page_start=page_start,
                            page_end=page_end,
                            section_path=section_path,
                            metadata={
                                "source_file": pdf_path.name,
                                "toc_level": toc_entry.level,
                                "chunk_method": "toc_based"
                            }
                        )
                        chunks.append(chunk)

                    if len(toc) >= 500 and ((i + 1) % 500 == 0 or i + 1 == len(toc)):
                        print(f"    ... processed {i + 1}/{len(toc)} TOC entries")
                
                print(f"  ✓ Created {len(chunks)} chunks based on TOC")
                
        except Exception as e:
            print(f"  ✗ Error chunking PDF with TOC: {e}")
        
        return chunks
    
    def _build_section_path(self, entry: TableOfContentsEntry, toc: List[TableOfContentsEntry], index: int) -> List[str]:
        """
        Build hierarchical section path for a TOC entry.
        
        Args:
            entry: Current TOC entry
            toc: Full TOC list
            index: Index of current entry
            
        Returns:
            List of section titles forming the path
        """
        path = []
        
        # Find parent entries by looking backward for higher-level items
        current_level = entry.level
        for i in range(index - 1, -1, -1):
            if toc[i].level < current_level:
                path.insert(0, toc[i].title)
                current_level = toc[i].level
                if current_level == 0:
                    break
        
        path.append(entry.title)
        return path
    
    def chunk_document_with_reasoning(self, file_path: Path) -> List[DocumentChunk]:
        """
        Chunk a document using AI reasoning when no TOC is available.
        
        Args:
            file_path: Path to document file
            
        Returns:
            List of document chunks
        """
        print(f"  → Using AI reasoning to chunk document...")
        
        # Read document content
        content = self._read_document_content(file_path)
        
        if not content.strip():
            print(f"  ✗ Empty or unreadable document")
            return []
        
        # Use AI to analyze document structure and suggest chunks
        chunks = self._analyze_and_chunk_with_ai(file_path, content)
        
        if chunks:
            print(f"  ✓ Created {len(chunks)} chunks using AI reasoning")
        else:
            print(f"  ℹ Falling back to simple chunking")
            chunks = self._simple_chunk(file_path, content)
        
        return chunks
    
    def _read_document_content(self, file_path: Path) -> str:
        """
        Read content from various document types.
        Falls back to OCR for scanned/image-based PDFs.
        
        Args:
            file_path: Path to document
            
        Returns:
            Document text content
        """
        try:
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n\n"
                    if content.strip():
                        return content
                # PyPDF2 returned empty text — try OCR
                return self._ocr_pdf(file_path)
            else:
                # Text-based files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read()
        except Exception as e:
            print(f"  ✗ Error reading document: {e}")
            return ""

    def _ocr_pdf(self, file_path: Path) -> str:
        """
        Extract text from a scanned/image-based PDF using OCR.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            OCR-extracted text content, or empty string if OCR unavailable
        """
        if not OCR_AVAILABLE:
            print(f"  ✗ Scanned PDF detected but OCR packages not installed (pytesseract, pdf2image)")
            return ""
        try:
            print(f"  → Scanned PDF detected, running OCR...")
            images = convert_from_path(str(file_path), dpi=300)
            ocr_text = ""
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image)
                ocr_text += page_text + "\n\n"
                if len(images) >= 10 and ((i + 1) % 5 == 0 or i + 1 == len(images)):
                    print(f"    ... OCR processed {i + 1}/{len(images)} pages")
            if ocr_text.strip():
                print(f"  ✓ OCR extracted {len(ocr_text)} characters from {len(images)} pages")
            else:
                print(f"  ✗ OCR produced no text from {len(images)} pages")
            return ocr_text
        except Exception as e:
            print(f"  ✗ OCR failed: {e}")
            return ""

    def _ocr_pdf_page(self, file_path: Path, page_num: int) -> str:
        """
        OCR a single page from a PDF.
        
        Args:
            file_path: Path to PDF file
            page_num: 0-based page number
            
        Returns:
            OCR-extracted text for that page, or empty string
        """
        if not OCR_AVAILABLE:
            return ""
        try:
            images = convert_from_path(
                str(file_path), dpi=300,
                first_page=page_num + 1, last_page=page_num + 1
            )
            if images:
                return pytesseract.image_to_string(images[0])
            return ""
        except Exception:
            return ""
    
    def _analyze_and_chunk_with_ai(self, file_path: Path, content: str) -> List[DocumentChunk]:
        """
        Use AI to analyze document structure and create intelligent chunks.
        
        Args:
            file_path: Path to document
            content: Document content
            
        Returns:
            List of document chunks
        """
        # Limit content size for API
        max_chars = get_config().get_max_content_for_analysis()
        content_sample = content[:max_chars]
        is_truncated = len(content) > max_chars
        
        content_display = f"DOCUMENT: {file_path.name}\n{f'[Content truncated - showing first {max_chars} characters]' if is_truncated else ''}\n\n{content_sample}"
        
        prompt = self.prompt_manager.format_prompt(
            "document_structure_analysis",
            content=content_display
        )

        try:
            response = self.client.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0.5,
                reasoning_effort=self.reasoning_effort
            )
            
            result = json.loads(response.choices[0].message.content)
            sections = result.get('sections', [])
            
            if not sections:
                return []
            
            # Create chunks based on AI analysis
            chunks = []
            for i, section in enumerate(sections):
                # Find content boundaries
                start_marker = section.get('start_marker', '')
                end_marker = section.get('end_marker', '')
                
                # Extract chunk content
                chunk_content = self._extract_chunk_content(
                    content, start_marker, end_marker, i == len(sections) - 1
                )
                
                if chunk_content.strip():
                    # Build section path
                    section_path = self._build_section_path_from_ai(sections, i)
                    
                    chunk = DocumentChunk(
                        chunk_id=f"{file_path.stem}_{i+1:03d}",
                        title=section.get('title', f"Section {i+1}"),
                        content=chunk_content.strip(),
                        section_path=section_path,
                        metadata={
                            "source_file": file_path.name,
                            "section_level": section.get('level', 0),
                            "summary": section.get('summary', ''),
                            "chunk_method": "ai_reasoning",
                            "document_type": result.get('document_type', 'unknown')
                        }
                    )
                    chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            print(f"  ✗ Error analyzing with AI: {e}")
            return []
    
    def _extract_chunk_content(self, full_content: str, start_marker: str, 
                               end_marker: str, is_last: bool) -> str:
        """
        Extract chunk content between markers.
        
        Args:
            full_content: Full document content
            start_marker: Start marker text
            end_marker: End marker text
            is_last: Whether this is the last chunk
            
        Returns:
            Extracted chunk content
        """
        try:
            # Find start position
            start_pos = full_content.find(start_marker)
            if start_pos == -1:
                # Try fuzzy matching
                start_pos = 0
            
            # Find end position
            if is_last or end_marker == "END_OF_DOCUMENT":
                end_pos = len(full_content)
            else:
                end_pos = full_content.find(end_marker, start_pos + len(start_marker))
                if end_pos == -1:
                    end_pos = len(full_content)
            
            return full_content[start_pos:end_pos]
            
        except Exception as e:
            print(f"  ✗ Error extracting chunk content: {e}")
            return ""
    
    def _build_section_path_from_ai(self, sections: List[Dict], index: int) -> List[str]:
        """
        Build hierarchical section path from AI-analyzed sections.
        
        Args:
            sections: List of section dictionaries from AI
            index: Current section index
            
        Returns:
            List of section titles forming the path
        """
        path = []
        current_level = sections[index].get('level', 0)
        
        # Find parent sections by looking backward
        for i in range(index - 1, -1, -1):
            if sections[i].get('level', 0) < current_level:
                path.insert(0, sections[i].get('title', f'Section {i+1}'))
                current_level = sections[i].get('level', 0)
                if current_level == 0:
                    break
        
        path.append(sections[index].get('title', f'Section {index+1}'))
        return path
    
    def _normalize_chunk_sizes(self, chunks: List[DocumentChunk], source_file: Path) -> List[DocumentChunk]:
        """
        Normalize chunk sizes by sub-chunking oversized chunks and merging undersized ones.
        
        Uses config values:
            - max_chunk_size: Chunks above this word count get sub-chunked
            - min_chunk_size: Chunks below this word count get merged with adjacent
            - chunk_size_target: Target word count for sub-chunks
        
        Args:
            chunks: List of document chunks from any chunker
            source_file: Original source file (for logging)
        
        Returns:
            Normalized list of DocumentChunk objects with word_count metadata
        """
        if not chunks:
            return chunks
        
        max_words = self.config.get_max_chunk_size()       # default 2000 words
        min_words = self.config.get_min_chunk_size()        # default 500 words
        target_words = self.config.get_chunk_size_target()  # default 1000 words
        
        # Convert target words to chars (~5 chars/word) for RecursiveCharacterTextSplitter
        target_chars = target_words * 5
        
        sub_chunk_count = 0
        merge_count = 0
        
        # --- Phase 1: Sub-chunk oversized chunks ---
        normalized = []
        for chunk in chunks:
            word_count = len(chunk.content.split())
            
            if word_count > max_words:
                # Split this chunk into sub-chunks
                sub_chunks = self._sub_chunk(chunk, target_chars, source_file)
                sub_chunk_count += len(sub_chunks) - 1  # net new chunks added
                normalized.extend(sub_chunks)
            else:
                normalized.append(chunk)
        
        # --- Phase 2: Merge undersized adjacent chunks within same parent section ---
        merged = []
        i = 0
        while i < len(normalized):
            chunk = normalized[i]
            word_count = len(chunk.content.split())
            
            # Only merge if undersized AND not already a sub-chunk AND has a next sibling
            if (word_count < min_words 
                and i + 1 < len(normalized)
                and not chunk.metadata.get('is_sub_chunk')
                and not normalized[i + 1].metadata.get('is_sub_chunk')):
                
                next_chunk = normalized[i + 1]
                next_word_count = len(next_chunk.content.split())
                
                # Check same parent section (all path elements except last match)
                same_parent = (chunk.section_path[:-1] == next_chunk.section_path[:-1]) if (chunk.section_path and next_chunk.section_path) else False
                
                # Merge if combined size is within max and same parent
                if same_parent and (word_count + next_word_count) <= max_words:
                    merged_content = chunk.content + "\n\n" + next_chunk.content
                    merged_title = f"{chunk.title} + {next_chunk.title}"
                    
                    merged_chunk = DocumentChunk(
                        chunk_id=chunk.chunk_id,
                        title=merged_title[:100],
                        content=merged_content,
                        page_start=chunk.page_start,
                        page_end=next_chunk.page_end or chunk.page_end,
                        section_path=chunk.section_path,
                        metadata={
                            **chunk.metadata,
                            "merged_with": next_chunk.chunk_id,
                            "merge_reason": f"both chunks under {min_words} words"
                        }
                    )
                    merged.append(merged_chunk)
                    merge_count += 1
                    i += 2  # Skip the next chunk (it's been merged)
                    continue
            
            merged.append(chunk)
            i += 1
        
        # --- Phase 3: Stamp word_count on every chunk ---
        for chunk in merged:
            chunk.metadata['word_count'] = len(chunk.content.split())
        
        # Log normalization results
        if sub_chunk_count > 0 or merge_count > 0:
            print(f"  \U0001f4cf Chunk normalization: {len(chunks)} \u2192 {len(merged)} chunks "
                  f"(sub-chunked {sub_chunk_count}, merged {merge_count})")
            word_counts = [c.metadata['word_count'] for c in merged]
            print(f"     Word range: {min(word_counts)}\u2013{max(word_counts)} "
                  f"(target: {min_words}\u2013{max_words}, ideal: {target_words})")
        else:
            print(f"  \u2713 All {len(merged)} chunks within size bounds ({min_words}\u2013{max_words} words)")
        
        return merged
    
    def _sub_chunk(self, chunk: DocumentChunk, target_chars: int, source_file: Path) -> List[DocumentChunk]:
        """
        Split an oversized chunk into sub-chunks preserving section hierarchy.
        
        Args:
            chunk: The oversized DocumentChunk
            target_chars: Target size in characters for each sub-chunk
            source_file: Source file for context
        
        Returns:
            List of sub-chunks with parent traceability metadata
        """
        content = chunk.content
        
        # Try langchain RecursiveCharacterTextSplitter first
        if LANGCHAIN_SUPPORT:
            try:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=target_chars,
                    chunk_overlap=get_config().get_chunk_overlap(),
                    separators=["\n\n", "\n", ". ", " ", ""]
                )
                parts = splitter.split_text(content)
            except Exception:
                parts = self._simple_text_split(content, target_chars)
        else:
            parts = self._simple_text_split(content, target_chars)
        
        if len(parts) <= 1:
            return [chunk]
        
        sub_chunks = []
        for idx, part in enumerate(parts):
            sub_id = f"{chunk.chunk_id}_sub_{idx + 1:03d}"
            sub_title = f"{chunk.title} (part {idx + 1}/{len(parts)})"
            
            sub_chunk = DocumentChunk(
                chunk_id=sub_id,
                title=sub_title,
                content=part.strip(),
                page_start=chunk.page_start,
                page_end=chunk.page_end,
                section_path=chunk.section_path,
                metadata={
                    **chunk.metadata,
                    "is_sub_chunk": True,
                    "parent_chunk_id": chunk.chunk_id,
                    "sub_chunk_index": idx + 1,
                    "total_sub_chunks": len(parts),
                    "parent_word_count": len(content.split())
                }
            )
            sub_chunks.append(sub_chunk)
        
        return sub_chunks
    
    def _simple_text_split(self, text: str, target_chars: int) -> List[str]:
        """
        Simple paragraph-aware text splitting fallback when langchain is unavailable.
        
        Args:
            text: Text to split
            target_chars: Target size per part
        
        Returns:
            List of text parts
        """
        paragraphs = text.split('\n\n')
        parts = []
        current = ""
        
        for para in paragraphs:
            if len(current) + len(para) + 2 > target_chars and current:
                parts.append(current)
                current = para
            else:
                current = current + "\n\n" + para if current else para
        
        if current:
            parts.append(current)
        
        return parts
    
    def _simple_chunk(self, file_path: Path, content: str, chunk_size: int = None) -> List[DocumentChunk]:
        """
        Simple fallback chunking by character count.
        
        Args:
            file_path: Path to document
            content: Document content
            chunk_size: Target size for each chunk
        
        Returns:
            List of document chunks
        """
        chunk_size = chunk_size or get_config().get_simple_chunk_size()
        chunks = []
        
        # Split into paragraphs first
        paragraphs = content.split('\n\n')
        
        current_chunk = ""
        chunk_num = 1
        
        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk.strip():
                    chunk = DocumentChunk(
                        chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                        title=f"Part {chunk_num}",
                        content=current_chunk.strip(),
                        section_path=[f"Part {chunk_num}"],
                        metadata={
                            "source_file": file_path.name,
                            "chunk_method": "simple"
                        }
                    )
                    chunks.append(chunk)
                    chunk_num += 1
                
                current_chunk = para + "\n\n"
        
        # Add last chunk
        if current_chunk.strip():
            chunk = DocumentChunk(
                chunk_id=f"{file_path.stem}_{chunk_num:03d}",
                title=f"Part {chunk_num}",
                content=current_chunk.strip(),
                section_path=[f"Part {chunk_num}"],
                metadata={
                    "source_file": file_path.name,
                    "chunk_method": "simple"
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def organize_chunks(self, chunks: List[DocumentChunk], source_file: Path, 
                       output_base: Path) -> Dict[str, Any]:
        """
        Organize chunks into folder structure.
        
        Args:
            chunks: List of document chunks
            source_file: Original source file
            output_base: Base output directory
            
        Returns:
            Dictionary with organization statistics
        """
        print(f"  → Organizing {len(chunks)} chunks into folder structure...")
        
        # Create folder for this document (sanitized to match naming rule)
        doc_folder = output_base / self._sanitize_folder_name(source_file.stem)
        doc_folder.mkdir(parents=True, exist_ok=True)
        
        # Save document metadata
        metadata = {
            "source_file": source_file.name,
            "total_chunks": len(chunks),
            "chunk_methods": list(set(c.metadata.get('chunk_method', 'unknown') for c in chunks)),
            "structure": []
        }
        
        # Organize chunks
        for chunk in chunks:
            # Create subfolder based on section path
            chunk_folder = doc_folder
            for section in chunk.section_path[:-1]:  # All but last item in path
                safe_section = self._sanitize_folder_name(section)
                chunk_folder = chunk_folder / safe_section
                chunk_folder.mkdir(parents=True, exist_ok=True)
            
            # Save chunk content
            chunk_filename = self._sanitize_folder_name(chunk.title) + ".txt"
            chunk_file = chunk_folder / chunk_filename
            
            with open(chunk_file, 'w', encoding='utf-8') as f:
                # Write header with metadata
                f.write(f"# {chunk.title}\n\n")
                f.write(f"**Source:** {chunk.metadata.get('source_file', 'Unknown')}\n")
                f.write(f"**Chunk ID:** {chunk.chunk_id}\n")
                
                if chunk.page_start:
                    f.write(f"**Pages:** {chunk.page_start}-{chunk.page_end}\n")
                
                f.write(f"**Section Path:** {' > '.join(chunk.section_path)}\n")
                
                if chunk.metadata.get('summary'):
                    f.write(f"**Summary:** {chunk.metadata['summary']}\n")
                
                f.write("\n---\n\n")
                f.write(chunk.content)
            
            # Add to metadata structure
            metadata["structure"].append({
                "chunk_id": chunk.chunk_id,
                "title": chunk.title,
                "path": str(chunk_file.relative_to(output_base)),
                "section_path": chunk.section_path,
                "metadata": chunk.metadata
            })
        
        # Save metadata file
        metadata_file = doc_folder / "_metadata.json"
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False)
        
        print(f"  ✓ Organized into: {doc_folder}")
        print(f"  ✓ Created {len(chunks)} chunk files")
        
        return metadata
    
    def _sanitize_folder_name(self, name: str) -> str:
        """
        Sanitize a string for use as folder/file name.

        Naming rule (matches compliance-files convention):
          - lowercase every character
          - replace any run of non-alphanumeric characters with a single '_'
          - strip leading/trailing underscores
          - cap length at 100 chars

        Args:
            name: Original name

        Returns:
            Sanitized name
        """
        sanitized = name.lower()
        sanitized = re.sub(r'[^a-z0-9]+', '_', sanitized)
        sanitized = sanitized.strip('_')[:100].strip('_')

        if not sanitized:
            sanitized = "untitled"

        return sanitized
    
    def process_knowledge_folder(self, input_path_str: str, 
                                 output_folder: str = "knowledge-files-organized",
                                 only_files: List[str] = None) -> Dict[str, Any]:
        """
        Process knowledge files (single file or folder).
        
        Args:
            input_path_str: Path to a file or folder containing knowledge files
            output_folder: Name of output folder (created in same parent)
            only_files: Optional list of specific filenames to process
            
        Returns:
            Dictionary with processing results
        """
        print(f"\n{'='*70}")
        print(f"KNOWLEDGE ORGANIZER AGENT")
        print(f"{'='*70}\n")
        
        # Setup paths
        input_path = Path(input_path_str)
        output_folder_path = Path(output_folder)
        
        # If output_folder is an absolute path or contains directories, use it directly
        # Otherwise treat it as a folder name relative to the input parent
        if output_folder_path.is_absolute() or '/' in output_folder:
            output_path = output_folder_path
        else:
            output_path = input_path.parent / output_folder
        
        output_path.mkdir(parents=True, exist_ok=True)
        
        if input_path.is_file():
            print(f"Input file: {input_path}")
        else:
            print(f"Input folder: {input_path}")
        print(f"Output folder: {output_path}\n")
        
        # Explore input (file or folder)
        print("Exploring knowledge files...")
        print(f"Supported formats: {', '.join(sorted(self.supported_extensions))}")
        files = self.explore_input(input_path_str, only_files=only_files)
        
        if not files:
            print("\n✗ No supported files found!")
            return {"status": "no_files", "processed": 0}
        
        print(f"\n{'─'*70}\n")
        
        # Process each file
        results = {
            "total_files": len(files),
            "processed": 0,
            "failed": 0,
            "total_chunks": 0,
            "chunker_tools_used": {},
            "chunk_word_counts": [],  # Track all chunk word counts for summary
            "sub_chunks_created": 0,
            "merges_performed": 0,
            "files": []
        }
        
        for i, file_path in enumerate(files, 1):
            print(f"[{i}/{len(files)}] Processing: {file_path.name}")
            print(f"{'─'*70}")
            
            try:
                chunks = []
                chunker_used = None
                
                # First, check if a specialized chunker tool can handle this file
                chunker = self.get_chunker_for_file(file_path)
                
                if chunker:
                    # Use specialized chunker tool
                    print(f"  → Using {chunker.name} tool...")
                    chunks = chunker.chunk(file_path)
                    chunker_used = chunker.name
                
                # PDF handling (built-in with TOC support)
                elif file_path.suffix.lower() == '.pdf':
                    print(f"  → Using PDFChunker (built-in)...")
                    print(f"  → Analyzing PDF structure...")
                    toc = self.extract_pdf_toc(file_path)
                    
                    if toc:
                        chunks = self.chunk_pdf_with_toc(file_path, toc)
                        chunker_used = "PDFChunker (TOC)"
                    else:
                        chunks = self.chunk_document_with_reasoning(file_path)
                        chunker_used = "PDFChunker (AI)"
                
                # Fallback: use reasoning-based chunking for text files
                elif not chunks:
                    print(f"  → Using TextChunker (AI reasoning)...")
                    chunks = self.chunk_document_with_reasoning(file_path)
                    chunker_used = "TextChunker"
                
                if chunks:
                    # Track pre-normalization count
                    pre_norm_count = len(chunks)
                    
                    # Normalize chunk sizes (sub-chunk oversized, merge undersized)
                    chunks = self._normalize_chunk_sizes(chunks, file_path)
                    
                    # Track normalization stats
                    results["sub_chunks_created"] += max(0, len(chunks) - pre_norm_count)
                    results["merges_performed"] += max(0, pre_norm_count - len(chunks))
                    
                    # Collect word counts for summary
                    for c in chunks:
                        wc = c.metadata.get('word_count', len(c.content.split()))
                        results["chunk_word_counts"].append(wc)
                    
                    # Organize chunks into folders
                    file_metadata = self.organize_chunks(chunks, file_path, output_path)
                    
                    results["processed"] += 1
                    results["total_chunks"] += len(chunks)
                    
                    # Track chunker tool usage
                    if chunker_used:
                        results["chunker_tools_used"][chunker_used] = results["chunker_tools_used"].get(chunker_used, 0) + 1
                    
                    results["files"].append({
                        "file": file_path.name,
                        "status": "success",
                        "chunks": len(chunks),
                        "method": chunks[0].metadata.get('chunk_method', 'unknown') if chunks else 'unknown',
                        "chunker_tool": chunker_used
                    })
                else:
                    print(f"  ✗ No chunks created")
                    results["failed"] += 1
                    results["files"].append({
                        "file": file_path.name,
                        "status": "failed",
                        "chunks": 0
                    })
                
            except Exception as e:
                print(f"  ✗ Error processing file: {e}")
                results["failed"] += 1
                results["files"].append({
                    "file": file_path.name,
                    "status": "error",
                    "error": str(e)
                })
            
            print()
            
            # Brief pause between files
            if i < len(files):
                time.sleep(1)
        
        # Save overall results
        results_file = output_path / "_processing_results.json"
        
        # Print summary
        print(f"{'='*70}")
        print(f"PROCESSING COMPLETE")
        print(f"{'='*70}")
        print(f"\nResults:")
        print(f"  Total files: {results['total_files']}")
        print(f"  Successfully processed: {results['processed']}")
        print(f"  Failed: {results['failed']}")
        print(f"  Total chunks created: {results['total_chunks']}")
        print(f"\nChunker Tools Used:")
        for tool, count in results["chunker_tools_used"].items():
            print(f"  - {tool}: {count} file(s)")
        
        # Chunk size distribution report
        wc_list = results["chunk_word_counts"]
        if wc_list:
            wc_sorted = sorted(wc_list)
            total_wc = len(wc_sorted)
            median_wc = wc_sorted[total_wc // 2]
            p10 = wc_sorted[int(total_wc * 0.1)] if total_wc >= 10 else wc_sorted[0]
            p90 = wc_sorted[int(total_wc * 0.9)] if total_wc >= 10 else wc_sorted[-1]
            under_500 = sum(1 for w in wc_sorted if w < 500)
            over_2000 = sum(1 for w in wc_sorted if w > 2000)
            
            print(f"\n📐 Chunk Size Distribution:")
            print(f"  Total chunks: {total_wc}")
            print(f"  Word count range: {wc_sorted[0]}–{wc_sorted[-1]}")
            print(f"  Average: {sum(wc_sorted) // total_wc} words")
            print(f"  Median: {median_wc} words")
            print(f"  P10–P90: {p10}–{p90} words")
            if results['sub_chunks_created'] > 0:
                print(f"  Sub-chunks created: {results['sub_chunks_created']}")
            if results['merges_performed'] > 0:
                print(f"  Merges performed: {results['merges_performed']}")
            if under_500 > 0:
                print(f"  ⚠ Under 500 words: {under_500} chunks")
            if over_2000 > 0:
                print(f"  ⚠ Over 2000 words: {over_2000} chunks")
        
        # Remove word_counts array from saved results (too verbose)
        results_to_save = {k: v for k, v in results.items() if k != 'chunk_word_counts'}
        results_to_save['chunk_size_stats'] = {
            'min_words': wc_sorted[0] if wc_list else 0,
            'max_words': wc_sorted[-1] if wc_list else 0,
            'avg_words': sum(wc_sorted) // len(wc_sorted) if wc_list else 0,
            'median_words': median_wc if wc_list else 0,
            'sub_chunks_created': results['sub_chunks_created'],
            'merges_performed': results['merges_performed']
        } if wc_list else {}
        
        print(f"\nOutput location: {output_path}")
        
        # Save results (without the verbose word_counts array)
        with open(results_file, 'w', encoding='utf-8') as f:
            json.dump(results_to_save, f, indent=2, ensure_ascii=False)
        
        print(f"Results saved to: {results_file}")
        print(f"\n{'='*70}\n")
        
        return results


def main():
    """
    Main execution function.
    """
    import sys
    
    # Load configuration
    sys.path.insert(0, str(Path(__file__).parent))
    from utils.config import get_config
    
    try:
        config = get_config()
        API_KEY = config.get_openai_api_key()
        REASONING_MODEL = config.get_reasoning_model()
    except (FileNotFoundError, ValueError) as e:
        print(f"❌ Error: {e}")
        print("\nSet OPENAI_API_KEY environment variable:")
        print("  export OPENAI_API_KEY='your-api-key-here'")
        sys.exit(1)
    
    # Get input folder from command line or use default
    if len(sys.argv) > 1:
        input_folder = sys.argv[1]
    else:
        input_folder = "knowledge-files"
        print(f"Using default input folder: {input_folder}")
        print(f"To specify a different folder, run: python {sys.argv[0]} <folder_path>\n")
    
    # Get output folder name (optional)
    output_folder = sys.argv[2] if len(sys.argv) > 2 else "knowledge-files-organized"
    
    # Parse --files argument if present (remaining args after first two positional)
    only_files = None
    if "--files" in sys.argv:
        files_idx = sys.argv.index("--files")
        only_files = sys.argv[files_idx + 1:]
        if only_files:
            print(f"\n📋 File filter active: processing only {len(only_files)} selected file(s)")
            for f in only_files:
                print(f"   • {f}")
    
    print("""
╔══════════════════════════════════════════════════════════════════════╗
║          KNOWLEDGE ORGANIZER AGENT                                   ║
║   Intelligent Document Chunking & Organization                       ║
╚══════════════════════════════════════════════════════════════════════╝
    """)
    REASONING_EFFORT = config.get_reasoning_effort()
    print(f"Model: {REASONING_MODEL}")
    print(f"  Reasoning Effort: {REASONING_EFFORT}")
    print()
    
    # Initialize agent
    agent = DocumentChunkingAgent(api_key=API_KEY, model=REASONING_MODEL, reasoning_effort=REASONING_EFFORT)
    
    # Process knowledge folder
    try:
        results = agent.process_knowledge_folder(input_folder, output_folder, only_files=only_files)
        
        if results["processed"] > 0:
            print("✓ Knowledge organization completed successfully!")
        else:
            print("⚠ No files were successfully processed")
            sys.exit(1)
            
    except Exception as e:
        print(f"\n❌ Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
