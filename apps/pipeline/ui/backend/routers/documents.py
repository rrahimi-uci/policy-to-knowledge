"""Documents router — list, upload, preview, delete compliance documents."""

import json
import mimetypes
import os
import re
import shutil
import zipfile
from pathlib import Path
from typing import List
from xml.etree import ElementTree as ET

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent

RAW_MEDIA_TYPES = {
    ".pdf": "application/pdf",
    ".txt": "text/plain; charset=utf-8",
    ".md": "text/markdown; charset=utf-8",
    ".csv": "text/csv; charset=utf-8",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
}

router = APIRouter(prefix="/api/documents", tags=["documents"])

SUPPORTED_EXT = {".pdf", ".txt", ".md", ".docx", ".csv", ".xlsx", ".pptx"}
SUPPORTED_DOMAINS = {"mortgage", "aml", "healthcare", "commercial_lending"}


def _safe_path(base: Path, *parts: str) -> Path:
    """Resolve path and ensure it stays within base (prevent traversal)."""
    base_resolved = base.resolve()
    resolved = (base / Path(*parts)).resolve()
    # Use is_relative_to to avoid false positives from str.startswith
    # e.g. /data/files vs /data/files_other
    if not resolved.is_relative_to(base_resolved):
        raise HTTPException(400, "Invalid path")
    return resolved


def _extract_pdf_text(filepath: Path, max_chars: int = 10000) -> str:
    """Extract text from a PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(filepath))
        parts = []
        total = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            parts.append(text)
            total += len(text)
            if total >= max_chars:
                break
        content = "\n\n".join(parts)
        return content[:max_chars] if len(content) > max_chars else content
    except Exception:
        return ""


def _truncate_preview(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _extract_docx_html(filepath: Path, max_chars: int = 200000) -> tuple[str, bool]:
    """Convert a .docx file to semantic HTML using mammoth, preserving headings,
    lists, tables, bold/italic and images. Falls back to empty string when
    mammoth is unavailable or conversion fails."""
    try:
        import mammoth  # type: ignore

        with open(filepath, "rb") as fh:
            # Map common Word styles to semantic HTML tags.
            style_map = """
            p[style-name='Title'] => h1.doc-title:fresh
            p[style-name='Subtitle'] => h2.doc-subtitle:fresh
            p[style-name='Heading 1'] => h1:fresh
            p[style-name='Heading 2'] => h2:fresh
            p[style-name='Heading 3'] => h3:fresh
            p[style-name='Heading 4'] => h4:fresh
            p[style-name='Heading 5'] => h5:fresh
            p[style-name='Heading 6'] => h6:fresh
            p[style-name='Quote'] => blockquote:fresh > p:fresh
            p[style-name='Intense Quote'] => blockquote.intense:fresh > p:fresh
            r[style-name='Strong'] => strong
            r[style-name='Emphasis'] => em
            """
            result = mammoth.convert_to_html(fh, style_map=style_map)
        html = (result.value or "").strip()
        if not html:
            return "", False
        return _truncate_preview(html, max_chars)
    except ModuleNotFoundError:
        return "", False
    except Exception:
        return "", False


def _extract_docx_text(filepath: Path, max_chars: int = 10000) -> tuple[str, bool]:
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(filepath))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = (getattr(paragraph.style, "name", "") or "").lower()
            if style_name.startswith("heading"):
                match = re.search(r"(\d+)", style_name)
                level = min(int(match.group(1)), 6) if match else 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return _truncate_preview("\n\n".join(parts), max_chars)
    except Exception:
        return "", False


def _extract_xlsx_preview(filepath: Path, max_chars: int = 10000) -> tuple[str, bool]:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(filename=str(filepath), read_only=True, data_only=True)
        sheets = []
        total_chars = 0
        truncated = False

        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value) for value in row[:12]]
                if not any(values):
                    continue
                rows.append(values)
                total_chars += sum(len(value) for value in values)
                if row_index >= 40 or total_chars >= max_chars:
                    truncated = True
                    break

            if rows:
                sheets.append({"name": sheet.title, "rows": rows})
            if total_chars >= max_chars:
                truncated = True
                break

        return json.dumps(sheets), truncated
    except Exception:
        return "[]", False


def _slide_sort_key(name: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", name)
    return int(match.group(1)) if match else 0


def _extract_xml_text_nodes(root: ET.Element) -> list[str]:
    texts: list[str] = []
    for node in root.iter():
        if node.tag.rsplit("}", 1)[-1] != "t":
            continue
        text = (node.text or "").strip()
        if text:
            texts.append(text)
    return texts


def _extract_pptx_preview(filepath: Path, max_chars: int = 10000) -> tuple[str, bool]:
    try:
        slides = []
        total_chars = 0
        truncated = False

        with zipfile.ZipFile(filepath) as archive:
            slide_names = sorted(
                [name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")],
                key=_slide_sort_key,
            )

            for slide_position, slide_name in enumerate(slide_names, start=1):
                slide_number = _slide_sort_key(slide_name) or slide_position
                texts: list[str] = []
                try:
                    root = ET.fromstring(archive.read(slide_name))
                    texts = _extract_xml_text_nodes(root)
                except Exception:
                    texts = []

                if texts:
                    title = texts[0]
                    bullets = texts[1:9]
                else:
                    title = f"Slide {slide_number}"
                    bullets = []

                slides.append({"title": title, "bullets": bullets})
                total_chars += len(title) + sum(len(text) for text in bullets)

                if total_chars >= max_chars or slide_position >= 20:
                    truncated = slide_position < len(slide_names) or len(texts) > 9 or total_chars >= max_chars
                    break

        return json.dumps(slides), truncated
    except Exception:
        return "[]", False


def _raw_media_type(filepath: Path) -> str:
    if _has_pdf_signature(filepath):
        return "application/pdf"
    ext = filepath.suffix.lower()
    return RAW_MEDIA_TYPES.get(ext) or mimetypes.guess_type(str(filepath))[0] or "application/octet-stream"


def _has_pdf_signature(filepath: Path) -> bool:
    try:
        with open(filepath, "rb") as fh:
            return fh.read(5) == b"%PDF-"
    except Exception:
        return False


def _source_dir() -> Path:
    return PROJECT_ROOT / "compliance-files"


def _folder_domain_path() -> Path:
    return _source_dir() / ".folder_domains.json"


def _infer_folder_domain(name: str) -> str:
    lower = name.lower()
    if (
        lower.startswith("p2k")
        or lower.startswith("fannie")
        or lower.startswith("freddie")
        or lower.startswith("freddies")
        or lower.startswith("fnma")
        or lower.startswith("fhlmc")
    ):
        return "mortgage"
    return ""


def _load_folder_domains() -> dict[str, str]:
    metadata_path = _folder_domain_path()
    if not metadata_path.exists():
        return {}
    try:
        data = json.loads(metadata_path.read_text())
    except Exception:
        return {}
    if not isinstance(data, dict):
        return {}
    return {
        str(folder): str(domain)
        for folder, domain in data.items()
        if isinstance(folder, str) and isinstance(domain, str) and domain in SUPPORTED_DOMAINS
    }


def _save_folder_domains(folder_domains: dict[str, str]) -> None:
    metadata_path = _folder_domain_path()
    metadata_path.parent.mkdir(parents=True, exist_ok=True)
    metadata_path.write_text(json.dumps(folder_domains, indent=2, sort_keys=True))


def _set_folder_domain(folder_name: str, domain: str | None) -> None:
    if not folder_name:
        return
    folder_domains = _load_folder_domains()
    if domain and domain in SUPPORTED_DOMAINS:
        folder_domains[folder_name] = domain
    else:
        folder_domains.pop(folder_name, None)
    _save_folder_domains(folder_domains)


def _folder_domain(folder_name: str) -> str:
    if not folder_name:
        return ""
    return _load_folder_domains().get(folder_name, _infer_folder_domain(folder_name))


def _normalize_path_part(part: str, *, keep_suffix: bool) -> str:
    """Lowercase and replace runs of non-alphanumerics with '_'.

    Allowed characters in the stem: ``a-z`` and ``0-9``. Every other
    character (including spaces, punctuation, uppercase letters and Unicode)
    is collapsed to a single underscore. Leading/trailing underscores are
    stripped.

    When ``keep_suffix`` is True, the final dotted suffix (e.g. ``.pdf``) is
    preserved (lowercased) so file extensions remain valid.
    """
    if keep_suffix and Path(part).suffixes:
        suffix = "".join(s.lower() for s in Path(part).suffixes)
        stem = part[: len(part) - len("".join(Path(part).suffixes))]
    else:
        suffix = ""
        stem = part
    stem = re.sub(r"[^a-z0-9]+", "_", stem.lower()).strip("_")
    if not stem:
        stem = "untitled"
    return stem + suffix


def _relative_upload_path(candidate: str | None, fallback_name: str) -> Path:
    """Normalize a relative upload path and reject traversal attempts.

    Each path segment is normalized so it consists only of ``a-z``, ``0-9``
    and ``_``. The final segment keeps its (lowercased) file extension.
    """
    raw_path = Path(candidate or fallback_name)
    parts = [part for part in raw_path.parts if part not in {"", "."}]
    if not parts:
        parts = [fallback_name]
    if any(part == ".." for part in parts):
        raise HTTPException(400, "Invalid upload path")
    normalized = [
        _normalize_path_part(p, keep_suffix=(i == len(parts) - 1))
        for i, p in enumerate(parts)
    ]
    return Path(*normalized)


def _document_payload(file_path: Path, root: Path) -> dict:
    """Build a stable response payload for a stored document."""
    relative_path = str(file_path.relative_to(root))
    return {
        "name": file_path.name,
        "relative_path": relative_path,
        "path": str(file_path.relative_to(PROJECT_ROOT)),
        "size": file_path.stat().st_size,
        "extension": file_path.suffix.lower(),
    }


@router.get("")
def list_documents(subdir: str = None):
    """List compliance documents, optionally within a subdirectory."""
    base = _source_dir()
    target = _safe_path(base, subdir) if subdir else base
    if not target.exists():
        return {"documents": [], "subdirectories": []}

    docs = []
    subdirs = []
    for item in sorted(target.iterdir()):
        if item.is_dir() and not item.name.startswith("."):
            file_count = sum(1 for f in item.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED_EXT)
            subdirs.append({
                "name": item.name,
                "file_count": file_count,
                "domain": _folder_domain(item.name) or None,
            })
        elif item.is_file() and item.suffix.lower() in SUPPORTED_EXT:
            docs.append(_document_payload(item, target))

    return {"documents": docs, "subdirectories": subdirs}


@router.get("/{subdir}/files")
def list_subdir_files(subdir: str):
    """List files inside a specific subdirectory."""
    target = _safe_path(_source_dir(), subdir)
    if not target.is_dir():
        raise HTTPException(404, f"Subdirectory '{subdir}' not found")

    docs = []
    for f in sorted(target.rglob("*")):
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXT:
            docs.append(_document_payload(f, target))
    return {"documents": docs, "subdirectory": subdir}


@router.get("/preview/{subdir}/{filename:path}")
def preview_document(subdir: str, filename: str, max_chars: int = 10000):
    """Return first N characters of a document, including PDF text extraction."""
    filepath = _safe_path(_source_dir(), subdir, filename)
    if not filepath.exists():
        raise HTTPException(404, "File not found")

    ext = filepath.suffix.lower()
    if _has_pdf_signature(filepath):
        content = _extract_pdf_text(filepath, max_chars)
        if content.strip():
            return {"filename": filename, "content": content, "truncated": len(content) >= max_chars, "type": "pdf_text"}
        return {"filename": filename, "content": "__PDF_EMBED__", "truncated": False, "type": "pdf_embed"}

    if ext == ".md":
        md_cap = max(max_chars, 200000)
        content = filepath.read_text(errors="replace")[:md_cap]
        return {"filename": filename, "content": content, "truncated": len(content) >= md_cap, "type": "markdown"}
    elif ext == ".txt":
        content = filepath.read_text(errors="replace")[:max_chars]
        return {"filename": filename, "content": content, "truncated": len(content) >= max_chars, "type": "text"}
    elif ext == ".csv":
        content = filepath.read_text(errors="replace")[:max_chars]
        return {"filename": filename, "content": content, "truncated": len(content) >= max_chars, "type": "csv"}
    elif ext == ".docx":
        html_content, html_truncated = _extract_docx_html(filepath)
        if html_content.strip():
            return {
                "filename": filename,
                "content": html_content,
                "truncated": html_truncated,
                "type": "docx_html",
            }
        content, truncated = _extract_docx_text(filepath, max_chars)
        if not content.strip():
            return {"filename": filename, "content": f"[Preview not available for {ext}]", "truncated": False, "type": "text"}
        return {"filename": filename, "content": content, "truncated": truncated, "type": "docx_text"}
    elif ext == ".xlsx":
        content, truncated = _extract_xlsx_preview(filepath, max_chars)
        if content == "[]":
            return {"filename": filename, "content": f"[Preview not available for {ext}]", "truncated": False, "type": "text"}
        return {"filename": filename, "content": content, "truncated": truncated, "type": "xlsx_sheets"}
    elif ext == ".pptx":
        content, truncated = _extract_pptx_preview(filepath, max_chars)
        if content == "[]":
            return {"filename": filename, "content": f"[Preview not available for {ext}]", "truncated": False, "type": "text"}
        return {"filename": filename, "content": content, "truncated": truncated, "type": "pptx_slides"}
    elif ext == ".pdf":
        content = _extract_pdf_text(filepath, max_chars)
        if content.strip():
            return {"filename": filename, "content": content, "truncated": len(content) >= max_chars, "type": "pdf_text"}
        return {"filename": filename, "content": "__PDF_EMBED__", "truncated": False, "type": "pdf_embed"}
    else:
        return {"filename": filename, "content": f"[Preview not available for {ext}]", "truncated": False}


@router.get("/raw/{subdir}/{filename:path}")
def serve_raw_file(subdir: str, filename: str):
    """Serve a raw file for inline viewing or direct download."""
    filepath = _safe_path(_source_dir(), subdir, filename)
    if not filepath.exists():
        raise HTTPException(404, "File not found")
    ext = filepath.suffix.lower()
    if ext not in SUPPORTED_EXT:
        raise HTTPException(400, f"Unsupported preview format: {ext}")
    return FileResponse(
        filepath,
        media_type=_raw_media_type(filepath),
        headers={"Content-Disposition": f"inline; filename=\"{filepath.name}\""},
    )


@router.post("/upload")
async def upload_documents(
    subdir: str | None = None,
    files: List[UploadFile] = File(...),
    relative_paths: List[str] = Form(default=[]),
    domain: str | None = Form(default=None),
):
    """Upload one or more compliance documents."""
    if domain and domain not in SUPPORTED_DOMAINS:
        raise HTTPException(400, "Invalid domain")

    base = _source_dir()
    target = _safe_path(base, subdir) if subdir else base
    target.mkdir(parents=True, exist_ok=True)

    uploaded = []
    created_folders = set()

    for index, f in enumerate(files):
        relative_path = _relative_upload_path(
            relative_paths[index] if index < len(relative_paths) else None,
            Path(f.filename).name,
        )
        ext = relative_path.suffix.lower()
        if ext not in SUPPORTED_EXT:
            continue

        destination_parts = [subdir] if subdir else []
        destination_parts.extend(relative_path.parts)
        dest = _safe_path(base, *destination_parts)
        dest.parent.mkdir(parents=True, exist_ok=True)
        with open(dest, "wb") as out:
            shutil.copyfileobj(f.file, out)
        if subdir:
            stored_root = subdir
            stored_relative = relative_path
        else:
            stored_root = relative_path.parts[0] if len(relative_path.parts) > 1 else ""
            stored_relative = relative_path
        if stored_root:
            created_folders.add(stored_root)
        uploaded.append({
            "name": dest.name,
            "relative_path": str(stored_relative),
            "folder": stored_root,
            "size": dest.stat().st_size,
        })

    primary_folder = subdir or (sorted(created_folders)[0] if len(created_folders) == 1 else None)
    if not subdir and domain:
        for folder_name in created_folders:
            _set_folder_domain(folder_name, domain)

    return {
        "uploaded": uploaded,
        "count": len(uploaded),
        "folders_created": sorted(created_folders),
        "primary_folder": primary_folder,
        "preserved_paths": bool(relative_paths),
        "domain": domain if domain in SUPPORTED_DOMAINS else (_folder_domain(primary_folder or "") or None),
    }


class CreateFolderRequest(BaseModel):
    name: str
    domain: str | None = None


@router.post("/folder")
def create_folder(body: CreateFolderRequest):
    """Create a new subdirectory inside compliance-files."""
    raw = body.name.strip()
    if not raw or "/" in raw or "\\" in raw or raw.startswith("."):
        raise HTTPException(400, "Invalid folder name")
    name = _normalize_path_part(raw, keep_suffix=False)
    if body.domain and body.domain not in SUPPORTED_DOMAINS:
        raise HTTPException(400, "Invalid domain")
    target = _safe_path(_source_dir(), name)
    if target.exists():
        raise HTTPException(409, f"Folder '{name}' already exists")
    target.mkdir(parents=True)
    if body.domain:
        _set_folder_domain(name, body.domain)
    return {"name": name, "domain": body.domain}


@router.delete("/folder/{subdir}")
def delete_folder(subdir: str):
    """Delete a subdirectory and all its contents."""
    target = _safe_path(_source_dir(), subdir)
    if not target.exists():
        raise HTTPException(404, f"Folder '{subdir}' not found")
    if not target.is_dir():
        raise HTTPException(400, "Not a directory")
    shutil.rmtree(target)
    _set_folder_domain(subdir, None)
    return {"deleted": subdir}


@router.delete("/file/{subdir}/{filename:path}")
def delete_file(subdir: str, filename: str):
    """Delete a specific file from a subdirectory."""
    filepath = _safe_path(_source_dir(), subdir, filename)
    if not filepath.exists():
        raise HTTPException(404, "File not found")
    if not filepath.is_file():
        raise HTTPException(400, "Not a file")
    filepath.unlink()
    return {"deleted": filename}
